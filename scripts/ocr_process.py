import os
import fitz # PyMuPDF
from paddleocr import PaddleOCR
from docx import Document
from docx.shared import Pt
import numpy as np
from PIL import Image

# Configuration
PDF_PATH = r"exam\Bản sao của 7-Vstep-Tests-B1-B2-C1-Full-Key.pdf"
# PDF_PATH = r"exam\test_sample.pdf" # For testing if needed
OUTPUT_DOCX = r"excel\7vstep_paddle_ocr.docx"

def main():
    if not os.path.exists(PDF_PATH):
        print(f"Error: PDF file not found at {PDF_PATH}")
        return

    # Initialize PaddleOCR
    # use_angle_cls=True helps with rotated text
    # lang='vi' for Vietnamese
    print("Initializing PaddleOCR...")
    
    # Fix for ConvertPirAttribute2RuntimeAttribute issue on Windows/CPU
    os.environ["FLAGS_use_mkldnn"] = "0"
    
    try:
        ocr = PaddleOCR(use_angle_cls=True, lang='vi', enable_mkldnn=False)
    except Exception as e:
        print(f"Error initializing PaddleOCR: {e}")
        # Try without specific args if it fails, or maybe just lang
        ocr = PaddleOCR(lang='vi', enable_mkldnn=False)

    import sys

    print(f"Opening PDF: {PDF_PATH}")
    doc = fitz.open(PDF_PATH)
    
    # Create valid output directory
    os.makedirs(os.path.dirname(OUTPUT_DOCX), exist_ok=True)
    
    start_page = 0
    if len(sys.argv) > 1:
        try:
            start_page = int(sys.argv[1])
        except ValueError:
            pass

    if start_page > 0 and os.path.exists(OUTPUT_DOCX):
        print(f"Resuming from page {start_page}. Loading existing docx...")
        document = Document(OUTPUT_DOCX)
    else:
        document = Document()
        # Set default font
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

    total_pages = len(doc)
    print(f"Total pages: {total_pages}. Starting from {start_page}.")

    # Process each page
    for page_num in range(start_page, total_pages):
        print(f"Processing page {page_num + 1}/{total_pages}...")
        
        # Periodic Save check (every 5 pages)
        if (page_num - start_page) > 0 and (page_num - start_page) % 5 == 0:
            print(f"Auto-saving progress to {OUTPUT_DOCX}...")
            document.save(OUTPUT_DOCX)

        page = doc.load_page(page_num)
        # ... (rest of loop)
        
        # Render page to image (high resolution for better OCR)
        zoom = 2.0 # 2x zoom for better clarity
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        
        # Convert fitz pixmap to PIL Image
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        img_np = np.array(img)
        
        # Run OCR
        # result = ocr.ocr(img_np, cls=True) 
        # API changed in new PaddleOCR/PaddleX: cls argument might be deprecated or implied
        try:
             result = ocr.ocr(img_np)
        except Exception as e:
             print(f"Error during OCR extraction: {e}")
             continue
        
        if not result or (isinstance(result, list) and not result[0]):
            print("No text found on this page.")
            document.add_page_break()
            continue

        # result structure check
        if isinstance(result, list) and result and isinstance(result[0], list) and result[0] and isinstance(result[0][0], list):
             # Old structure: [[[[x,y],..], (text, score)], ...]
             lines = result[0]
        elif isinstance(result, list) and result and isinstance(result[0], dict):
             # New structure maybe?
             # For now assume mostly compatible or flatten
             lines = result
        else:
             print(f"Unexpected result format: {type(result)}")
             # simplistic fallback
             lines = result if isinstance(result, list) else []

        # Validate lines structure [box, (text, info)]
        # If it's the new PaddleX return, it might be different.
        # Let's hope for the best or print structure if failed.
        
        # Simple post-processing and adding to docx
        # We aim to reconstruct paragraphs roughly.
        
        current_paragraph_text = []
        
        # Heuristics for paragraph reconstruction could be complex. 
        # For now, we will just treat each OCR line as a line in Word, 
        # but check for clear headers.
        
        for line in lines:
            if isinstance(line, dict):
                # New PaddleX/OCR structure?
                # Likely keys: 'points' or 'box', 'text' or 'rec_text', 'score' or 'rec_score'
                # Print keys once to debug if needed, but lets guess common ones
                box = line.get('points', line.get('box', []))
                text = line.get('transcription', line.get('text', line.get('rec_text', '')))
                confidence = line.get('score', line.get('rec_score', 0.0))
            elif isinstance(line, list):
                 # List structure [box, (text, score)]
                 if len(line) >= 2:
                     box = line[0]
                     text_obj = line[1]
                     if isinstance(text_obj, (list, tuple)) and len(text_obj) >= 2:
                         text = text_obj[0]
                         confidence = text_obj[1]
                     else:
                         text = str(text_obj)
                         confidence = 1.0 # Unknown
                 else:
                     continue
            else:
                continue

            # Basic cleanup
            if not isinstance(text, str):
                text = str(text)
            
            # Skip low confidence garbage? Maybe < 0.5
            if confidence < 0.5:
                continue
                
            # Check if it looks like a header
            # VSTEP structure often has "TEST 1", "SECTION 1", "PART 1"
            upper_text = text.upper()
            is_header = False
            if upper_text.startswith("TEST") and len(text) < 20:
                is_header = True
            elif upper_text.startswith("SECTION") or upper_text.startswith("PART"):
                is_header = True
            
            p = document.add_paragraph()
            if is_header:
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(14)
            else:
                p.add_run(text)
        
        # Add page break between pages to preserve structure
        document.add_page_break()

    print(f"Saving to {OUTPUT_DOCX}...")
    document.save(OUTPUT_DOCX)
    print("Done!")

if __name__ == "__main__":
    main()
